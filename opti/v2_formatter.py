import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Pt, RGBColor
from fasils import TYPE_A_XML, TYPE_B_XML
import re
import sys
import os
import requests
import win32com.client
import time

# --- Configuration ---
N8N_WEBHOOK_URL = "http://localhost:5679/webhook/0b2e5c27-0dc5-4512-8fcb-cbec9ba785fa"

def apply_formatting(run, font_name, size, bold=False, color=None):
    """
    Surgically builds the Run Properties (rPr) to ensure both standard 
    and Complex Script (Arabic) properties are set in the correct XML order.
    """
    rPr = run._element.get_or_add_rPr()
    
    # Clear existing properties to ensure strict ordering
    for child in rPr.getchildren():
        rPr.remove(child)
        
    # 1. Fonts (Standard and Complex Script)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    
    # 2. Bold (Standard and Complex Script)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
        
    # 3. Color
    if color:
        c = OxmlElement('w:color')
        hex_color = "%02x%02x%02x" % (color[0], color[1], color[2])
        c.set(qn('w:val'), hex_color)
        rPr.append(c)
        
    # 4. Size (Standard and Complex Script) — value is in half-points
    sz_val = str(int(size * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    
    # 5. Right-to-Left marker (Crucial for Arabic punctuation positioning)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def insert_fasil(doc, xml_string):
    para = doc.add_paragraph()
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    element = parse_xml(xml_string)
    para._p.append(element)

def apply_rtl_justify(para):
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def process_paragraph_text(para, text, base_font, base_size, base_color=None):
    base_phrases = [
        "صلى الله عليه وسلم",
        "رحمه الله تعالى",
        "حفظه الله تعالى",
        "سبحانه وتعالى",
        "رضي الله عنها",
        "رضي الله عنه",
        "رضي الله عنهم",
        "رضي الله عنهما",
        "رحمه الله",
        "حفظه الله",
        "عز وجل",
        "تعالى"
    ]
    
    # Generate diacritics-insensitive regex patterns for each phrase
    diacritics = r'[\u064B-\u0652]'
    honorifics = []
    for phrase in base_phrases:
        clean = re.sub(diacritics, '', phrase)
        pattern_parts = []
        for char in clean:
            if char.isspace():
                pattern_parts.append(r'\s+')
            else:
                pattern_parts.append(re.escape(char) + r'[\u064B-\u0652]*')
        honorifics.append(''.join(pattern_parts))
        
    honorific_subpattern = '|'.join(honorifics)
    pattern = rf'(<hadith>.*?</hadith>|<quran>.*?</quran>|<strong>.*?</strong>|\[.*?\]|{honorific_subpattern})'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part: continue
        
        run = para.add_run()
        
        is_honorific = re.match(rf'^({honorific_subpattern})$', part)
        
        if part.startswith('<hadith>') and part.endswith('</hadith>'):
            content = part.replace('<hadith>', '').replace('</hadith>', '')
            is_bold = False
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
                is_bold = True
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=is_bold, color=base_color)
            
        elif part.startswith('<quran>') and part.endswith('</quran>'):
            content = part.replace('<quran>', '').replace('</quran>', '')
            # If the verse starts and ends with decorative brackets
            if content.startswith('﴿') and content.endswith('﴾'):
                # 1. Opening bracket in Lotus font
                run.text = '﴿'
                apply_formatting(run, base_font, 18, color=RGBColor(0x20, 0x21, 0x22))
                # 2. Middle text in Uthmanic font
                run_mid = para.add_run(content[1:-1])
                apply_formatting(run_mid, 'KFGQPC Uthmanic Script HAFS', 18, color=RGBColor(0x20, 0x21, 0x22))
                # 3. Closing bracket in Lotus font
                run_end = para.add_run('﴾')
                apply_formatting(run_end, base_font, 18, color=RGBColor(0x20, 0x21, 0x22))
            else:
                run.text = content
                apply_formatting(run, 'KFGQPC Uthmanic Script HAFS', 18, color=RGBColor(0x20, 0x21, 0x22))
            
        elif part.startswith('<strong>') and part.endswith('</strong>'):
            content = part.replace('<strong>', '').replace('</strong>', '')
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=True, color=base_color)
            
        elif part.startswith('[') and part.endswith(']'):
            # Plain brackets = book titles or scholar quotes (NOT bold)
            content = part[1:-1]
            run.text = f">{content}<"
            apply_formatting(run, base_font, base_size, bold=False, color=base_color)
            
        elif is_honorific:
            run.text = part
            if 'Jameel Noori Nastaleeq' in base_font:
                apply_formatting(run, base_font, base_size, bold=False, color=base_color)
            else:
                apply_formatting(run, base_font, base_size, bold=False, color=RGBColor(0xFF, 0x00, 0x66))
            
        else:
            if 'Jameel Noori Nastaleeq' in base_font:
                run.text = part
                apply_formatting(run, base_font, base_size, bold=False, color=base_color)
            else:
                arabic_chars = r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFC]'
                sub_pattern = rf'(?<!{arabic_chars})(أَيْ|أي)(?!{arabic_chars})|([\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFCa-zA-Z0-9]+:)'
                sub_parts = re.split(sub_pattern, part)
                first = True
                for sub_part in sub_parts:
                    if sub_part is None: continue
                    if not sub_part: continue
                    if first:
                        current_run = run
                        first = False
                    else:
                        current_run = para.add_run()
                    current_run.text = sub_part
                    
                    is_ay = re.match(rf'^(?:أَيْ|أي|أَي)$', sub_part)
                    is_word_colon = re.match(r'^[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFCa-zA-Z0-9]+:$', sub_part)
                    
                    if is_ay or is_word_colon:
                        apply_formatting(current_run, base_font, base_size, bold=True, color=base_color)
                    else:
                        apply_formatting(current_run, base_font, base_size, bold=False, color=base_color)

def convert_to_pdf(docx_path, pdf_path):
    """Converts a Word document to PDF using win32com."""
    word = None
    wdoc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        
        wdoc = word.Documents.Open(abs_docx)
        # 17 = wdExportFormatPDF
        wdoc.ExportAsFixedFormat(abs_pdf, 17)
        return True
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return False
    finally:
        if wdoc: wdoc.Close(False)
        if word: word.Quit()

def send_to_n8n(docx_path, pdf_path):
    """Sends DOCX and PDF to n8n webhook."""
    try:
        files = {
            'docx': (os.path.basename(docx_path), open(docx_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'pdf': (os.path.basename(pdf_path), open(pdf_path, 'rb'), 'application/pdf'),
        }
        print(f"Sending to n8n: {os.path.basename(docx_path)}")
        resp = requests.post(N8N_WEBHOOK_URL, files=files)
        for f in files.values(): f[1].close()
        return resp.status_code in [200, 201]
    except Exception as e:
        print(f"n8n Upload Error: {e}")
        return False

def format_document(markdown_content, template_path, output_docx_path, is_file=True, page1_title=None, page2_title=None):
    """
    markdown_content: either a file path or a string of markdown.
    is_file: True if markdown_content is a path, False if it's raw text.
    page1_title: Book name for page 1 header replacement
    page2_title: Book name for page 2 header replacement
    """
    doc = docx.Document(template_path)
    
    # Fresh start
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # Replace book names in template headers if provided
    # if page1_title or page2_title:
    #     for section in doc.sections:
    #         if page1_title:
    #             section.different_first_page_header_footer = True
    #             p1 = section.first_page_header.paragraphs[0] if section.first_page_header.paragraphs else section.first_page_header.add_paragraph()
    #             p1.text = f"\t          {page1_title}"
            
    #         title_p2 = page2_title if page2_title else page1_title
    #         if title_p2:
    #             for para in section.header.paragraphs:
    #                 replaced = False
    #                 for run in para.runs:
    #                     if "كتاب فتح المجيد" in run.text:
    #                         run.text = run.text.replace("كتاب فتح المجيد", title_p2)
    #                         replaced = True
    #                     elif "[PAGE1_BOOK_NAME]" in run.text:
    #                         run.text = run.text.replace("[PAGE1_BOOK_NAME]", title_p2)
    #                         replaced = True
    #                     elif "[PAGE2_BOOK_NAME]" in run.text:
    #                         run.text = run.text.replace("[PAGE2_BOOK_NAME]", title_p2)
    #                         replaced = True
    #                 if not replaced:
    #                     para.text = f"\t          {title_p2}"
                        
    #         # Also check XML for elements
    #         header_xml = section._element
    #         for t_elem in header_xml.iter():
    #             if t_elem.tag.endswith('}t') and t_elem.text:
    #                 if page1_title and "[PAGE1_BOOK_NAME]" in t_elem.text:
    #                     t_elem.text = t_elem.text.replace("[PAGE1_BOOK_NAME]", page1_title)
    #                 if page2_title and "[PAGE2_BOOK_NAME]" in t_elem.text:
    #                     t_elem.text = t_elem.text.replace("[PAGE2_BOOK_NAME]", page2_title)
                        
    if is_file:
        with open(markdown_content, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    else:
        lines = markdown_content.splitlines()
        
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Convert double quotes to parentheses
        line = re.sub(r'["“]([^"“”]*?)["”]', r'(\1)', line)
            
        if line == '<matn>' or line == '</matn>': continue
            
        if line.startswith('<speaker>'):
            speaker_text = line.replace('<speaker>', '').replace('</speaker>', '')
            if 'قال الشارح' in speaker_text or 'قال المصنف' in speaker_text:
                if len(doc.paragraphs) > 0:
                    fasil_xml = TYPE_A_XML if 'قال الشارح' in speaker_text else TYPE_B_XML
                    insert_fasil(doc, fasil_xml)
            
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            process_paragraph_text(para, speaker_text, 'Jameel Noori Nastaleeq', 45, RGBColor(0xC0, 0x00, 0x00))
            continue
            
        if line.startswith('- '):
            line = line[2:]
            para = doc.add_paragraph(style='List Paragraph')
            apply_rtl_justify(para)
            bullet_run = para.add_run("○ ")
            apply_formatting(bullet_run, 'AAAGoldenLotus Stg1_Ver1', 18, bold=True, color=RGBColor(0x00, 0xB0, 0x50))
            process_paragraph_text(para, line, 'AAAGoldenLotus Stg1_Ver1', 18)
            continue
            
        para = doc.add_paragraph()
        apply_rtl_justify(para)
        process_paragraph_text(para, line, 'AAAGoldenLotus Stg1_Ver1', 18)
    
    # Final Green Symbol
    insert_fasil(doc, TYPE_A_XML)
    
    doc.save(output_docx_path)
    print(f"Saved DOCX: {output_docx_path}")
    
    # PDF Conversion
    pdf_path = output_docx_path.replace(".docx", ".pdf")
    if convert_to_pdf(output_docx_path, pdf_path):
        print(f"Saved PDF: {pdf_path}")
        # Send to n8n
        if send_to_n8n(output_docx_path, pdf_path):
            print("Successfully sent to n8n/Telegram!")
    
    return output_docx_path, pdf_path

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(base_dir)
    md_path = os.path.join(base_dir, "sample.md")
    template_path = os.path.join(base_dir, "template.docx")
    output_path = os.path.join(base_dir, "output.docx")
    format_document(md_path, template_path, output_path)
