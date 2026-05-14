import sys
import os
import time
import re
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import win32com.client
import win32clipboard
import win32con
import requests

# لضمان عمل الطباعة مع الحروف العربية في شاشة الأوامر
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

# ==========================================
# الإعدادات - نفس إعدادات مشروعك
# ==========================================
FONT_NAME = "AAAGoldenLotus Stg1_Ver1"
FONT_SIZE = 18
RED_HEX = "FF0000"
BLUE_HEX = "FF0066"
OUTPUT_DIR = r"C:\Users\saif_\Desktop\downs\حاليًا\يومي\Lectures\ANtiGrav"
N8N_WEBHOOK_URL = "https://atomicphoenix1.app.n8n.cloud/webhook/0b2e5c27-0dc5-4512-8fcb-cbec9ba785fa"
N8N_MP3_WEBHOOK_URL = "https://atomicphoenix1.app.n8n.cloud/webhook/6973bd50-c567-4e5b-8c56-51d59318dadd"

BLUE_PHRASES = [
    "صَلَّى اللَّهُ عَلَيْهِ وَسَلَّمَ",
    "رَضِيَ اللَّهُ عَنْهَا",
    "تَعَالَى",
    "رَضِيَ اللَّهُ عَنْهُ"
]

BLUE_RE = re.compile(f"({'|'.join(map(re.escape, BLUE_PHRASES))})")

def get_html_from_clipboard():
    """يسحب محتوى HTML من الحافظة"""
    try:
        win32clipboard.OpenClipboard()
        # تنسيق HTML المسجل في ويندوز
        try:
            format_id = win32clipboard.RegisterClipboardFormat("HTML Format")
        except:
            return None
            
        if win32clipboard.IsClipboardFormatAvailable(format_id):
            data = win32clipboard.GetClipboardData(format_id)
            if isinstance(data, bytes):
                # HTML في ويندوز يكون له رأس (Header) يحتوي على المسارات
                # سنحاول فك التشفير بصيغة utf-8
                return data.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Error reading clipboard: {e}")
    finally:
        try:
            win32clipboard.CloseClipboard()
        except:
            pass
    return None

def parse_html_to_segments(html_content):
    """
    يحول الـ HTML إلى قائمة من القطع (Segments) مع تحديد هل هي مائلة أم لا.
    """
    # تنظيف رأسية HTML في ويندوز
    if "StartFragment" in html_content:
        fragment_match = re.search(r'<!--StartFragment-->(.*)<!--EndFragment-->', html_content, re.DOTALL)
        if fragment_match:
            html_content = fragment_match.group(1)
    else:
        # إذا لم توجد العلامات، نحاول حذف الرأسية يدوياً إذا وجدت
        html_content = re.sub(r'^Version:.*?\r?\nStartHTML:.*?\r?\n', '', html_content, flags=re.DOTALL | re.IGNORECASE)

    soup = BeautifulSoup(html_content, 'html.parser')
    segments = []

    def is_italic(element):
        if element.name in ['i', 'em']:
            return True
        style = element.get('style', '').lower()
        if 'font-style: italic' in style or 'font-style:italic' in style:
            return True
        return False

    def process_element(element, inherited_italic=False):
        if element.name is None: # نص خام
            text = str(element)
            if text.strip():
                if segments and segments[-1].get('italic') == inherited_italic and 'type' not in segments[-1]:
                    segments[-1]['text'] += text
                else:
                    segments.append({'text': text, 'italic': inherited_italic})
            return

        current_italic = inherited_italic or is_italic(element)
        
        for child in element.children:
            if child.name == 'br':
                segments.append({'text': '\n', 'italic': False, 'type': 'br'})
            elif child.name in ['p', 'div', 'li']:
                process_element(child, current_italic)
                segments.append({'text': '\n', 'italic': False, 'type': 'block'})
            else:
                process_element(child, current_italic)

    # تنظيف السوب من الفراغات المزعجة في البداية
    process_element(soup)
    return segments

def heal_and_format_segments(segments):
    """يعالج انكسار الأسطر ويجهز النص النهائي للوورد"""
    final_blocks = []
    current_block = []

    for seg in segments:
        text = seg['text']
        is_italic = seg.get('italic', False)
        seg_type = seg.get('type', '')

        if seg_type == 'block':
            if current_block:
                final_blocks.append(current_block)
                current_block = []
            continue
            
        if seg_type == 'br':
            # دمج الأسطر المكسورة في العربي:
            # إذا كان السطر السابق لا ينتهي بنقطة، نحول الـ br لمسافة
            if current_block:
                last_text = current_block[-1]['text']
                if not last_text.strip().endswith(('.', ':', '؟', '!', '﴿', '﴾')):
                    current_block[-1]['text'] += " "
                else:
                    # سطر جديد حقيقي (فقرة جديدة داخلية)
                    final_blocks.append(current_block)
                    current_block = []
            continue

        # دمج النصوص المتتالية بنفس التنسيق
        if current_block and current_block[-1]['italic'] == is_italic:
            current_block[-1]['text'] += text
        else:
            current_block.append({'text': text, 'italic': is_italic})

    if current_block:
        final_blocks.append(current_block)
        
    return final_blocks

def apply_complex_style(run, is_italic=False, color_hex=None):
    """تطبيق التنسيق الكامل (خط 18، عريض، لوتس، واتجاه عربي)"""
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.bold = False
    
    # تحديد اللون: الأزرق له الأولوية، ثم الأحمر للمائل
    final_color = color_hex
    if not final_color and is_italic:
        final_color = RED_HEX

    if is_italic:
        run.font.italic = True
    
    if final_color:
        r, g, b = int(final_color[:2], 16), int(final_color[2:4], 16), int(final_color[4:], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), FONT_NAME)
    
    szCs = docx.oxml.shared.OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(FONT_SIZE * 2))
    rPr.append(szCs)
    
    bCs = docx.oxml.shared.OxmlElement('w:bCs')
    bCs.set(qn('w:val'), '1')
    rPr.append(bCs)
    
    if is_italic:
        iCs = docx.oxml.shared.OxmlElement('w:iCs')
        iCs.set(qn('w:val'), '1')
        rPr.append(iCs)
    
    if final_color:
        color_elem = docx.oxml.shared.OxmlElement('w:color')
        color_elem.set(qn('w:val'), final_color)
        rPr.append(color_elem)

def send_to_n8n(docx_path, pdf_path):
    """إرسال الملفات إلى n8n عبر Webhook"""
    # تحقق إذا كان الرابط لا يزال هو الرابط الافتراضي أو فارغاً
    if not N8N_WEBHOOK_URL or "your-webhook-url" in N8N_WEBHOOK_URL:
        print("[-] لم يتم تكوين رابط n8n أو الرابط غير صالح.")
        return

    try:
        # 1. إرسال المستندات (Word & PDF)
        files_docs = {
            'docx': (os.path.basename(docx_path), open(docx_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'pdf': (os.path.basename(pdf_path), open(pdf_path, 'rb'), 'application/pdf'),
        }
        
        print(f"[+] Sending Word and PDF documents to n8n...")
        resp1 = requests.post(N8N_WEBHOOK_URL, files=files_docs)
        if resp1.status_code in [200, 201]:
            print("[!] Documents sent successfully.")
        else:
            print(f"[-] Failed to send documents: {resp1.status_code}")
        
        for f in files_docs.values(): f[1].close()
            
    except Exception as e:
        print(f"[-] Error during n8n upload: {e}")

def get_next_output_path():
    """تحديد المسار التالي المتاح (ai_studio_output_1, 2, ...)"""
    base_name = "ai_studio_output"
    index = 1
    while True:
        path = os.path.join(OUTPUT_DIR, f"{base_name}_{index}.docx")
        if not os.path.exists(path):
            return path
        index += 1

def create_word_doc(blocks):
    doc = docx.Document()
    
    for block in blocks:
        p = doc.add_paragraph()
        for seg in block:
            # 1. استبدال الأقواس
            text = seg['text'].replace('[', '>').replace(']', '<').replace('\n', ' ').strip()
            if not text: continue
            
            # 2. البحث عن العبارات الزرقاء وتقسيم النص
            parts = BLUE_RE.split(text)
            for part in parts:
                if not part: continue
                is_blue = part in BLUE_PHRASES
                
                run = p.add_run(part + " ")
                # نمرر لوناً أزرق إذا كانت العبارة زرقاء، وإلا سيأخذ الأحمر إذا كان مائلاً
                apply_complex_style(run, is_italic=seg['italic'], color_hex=BLUE_HEX if is_blue else None)

    # تحديد اسم الملف الجديد (تلقائياً: 1، 2، 3...)
    final_output_path = get_next_output_path()
    temp_path = final_output_path.replace(".docx", "_temp.docx")
    
    # محاولة حذف الملف المؤقت القديم إذا كان موجوداً
    try:
        if os.path.exists(temp_path):
            os.remove(temp_path)
    except PermissionError:
        print(f"[-] خطأ: الملف المؤقت {os.path.basename(temp_path)} مفتوح في برنامج آخر. يرجى إغلاق الوورد.")
        return False

    try:
        doc.save(temp_path)
    except PermissionError:
        print(f"[-] خطأ: لا يمكن حفظ الملف المؤقت. يرجى إغلاق أي مستندات وورد مفتوحة.")
        return False
    
    word = None
    wdoc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        # فتح الملف المؤقت كمسار مطلق
        abs_temp_path = os.path.abspath(temp_path)
        abs_output_path = os.path.abspath(final_output_path)
        
        wdoc = word.Documents.Open(abs_temp_path)
        wdoc.Content.ParagraphFormat.ReadingOrder = 1 # RTL
        wdoc.Content.ParagraphFormat.Alignment = 2    # Right
        
        # محاولة الحفظ في الملف النهائي
        wdoc.SaveAs(abs_output_path)
        
        # التصدير بصيغة PDF
        pdf_output_path = abs_output_path.replace(".docx", ".pdf")
        wdoc.ExportAsFixedFormat(pdf_output_path, 17) # 17 = wdExportFormatPDF
        
        print(f"[!] تم حفظ الملفات بنجاح: {os.path.basename(final_output_path)} و PDF")
        
        # إرسال الملفات إلى n8n
        send_to_n8n(abs_output_path, pdf_output_path)
        
        return True
    except Exception as e:
        print(f"[-] حدث خطأ أثناء تنسيق الوورد: {e}")
        return False
    finally:
        if wdoc:
            try: wdoc.Close(False) # إغلاق بدون حفظ التغييرات على المؤقت
            except: pass
        if word:
            try: word.Quit()
            except: pass
        
        # محاولة أخيرة لحذف المؤقت
        time.sleep(0.5)
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except:
            pass

def monitor_clipboard():
    print("=== AI Studio Smart Formatter ===")
    print("البرنامج يعمل الآن... قم بنسخ النص (Highlight & Copy) من AI Studio.")
    print("سيتم تلقائياً حفظ النص المنسق في ملف وورد.")
    print("اضغط Ctrl+C في هذه الشاشة للتوقف.")
    
    # تهيئة الحافظة بالمحتوى الحالي لمنع تكرار معالجة آخر شيء تم نسخه قبل تشغيل البرنامج
    last_html = get_html_from_clipboard() or ""
    
    try:
        while True:
            html = get_html_from_clipboard()
            if html and html != last_html:
                print("\n[+] تم رصد نص جديد في الحافظة...")
                last_html = html
                
                segments = parse_html_to_segments(html)
                if not segments:
                    print("[-] لم يتم العثور على محتوى قابل للتحليل في الـ HTML.")
                    continue
                    
                blocks = heal_and_format_segments(segments)
                if not blocks:
                    print("[-] بعد معالجة النص، لم يتبق أي محتوى لحفظه.")
                    continue

                print(f"[+] تم استخراج {len(blocks)} فقرة. جاري الحفظ...")
                if create_word_doc(blocks):
                    # إصدار صوت تنبيه بسيط
                    import winsound
                    winsound.Beep(1000, 200)
                else:
                    print("[-] فشل حفظ الملف.")
            
            time.sleep(1.5)
    except KeyboardInterrupt:
        print("\nتم إيقاف البرنامج.")

if __name__ == "__main__":
    # التأكد من وجود المجلد
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    monitor_clipboard()
