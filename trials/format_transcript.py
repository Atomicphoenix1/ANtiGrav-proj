import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import win32com.client
import os
import re

# ==========================================
# 1. الإعدادات والخيارات
# ==========================================
INPUT_FILE = r"C:\Users\saif_\Desktop\downs\حاليًا\يومي\Lectures\ANtiGrav\mid.docx"
OUTPUT_FILE = r"C:\Users\saif_\Desktop\downs\حاليًا\يومي\Lectures\ANtiGrav\تفريغ.docx"

# خيار تلوين ما بين الأقواس (غير المائلة)
COLOR_BRACKETS = False  # True للأحمر، False للأسود

FONT_NAME = "AAAGoldenLotus Stg1_Ver1"
FONT_SIZE = 18
RED_HEX = "FF0000"

def heal_broken_lines(doc):
    """
    يعالج مشكلة انكسار الأسطر مع الحفاظ على الأقسام (Sections)
    عن طريق دمج الأسطر المتتالية التي لا يفصل بينها سطر فارغ.
    """
    sections = []
    current_chunk = []

    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            # إذا وجدنا سطر فارغ، ندمج ما سبق ونعتبره قسماً مستقلاً
            if current_chunk:
                sections.append(" ".join(current_chunk))
                current_chunk = []
            continue
        
        # إضافة السطر الحالي للمجموعة ليتم دمجها بمسافة واحدة
        current_chunk.append(text)
            
    if current_chunk:
        sections.append(" ".join(current_chunk))
        
    return sections

def apply_complex_style(run, is_red=False, is_italic=False):
    """تطبيق التنسيق الكامل (خط 18، عريض، لوتس، واتجاه عربي)"""
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.bold = True
    if is_red: run.font.color.rgb = RGBColor(255, 0, 0)
    if is_italic: run.font.italic = True

    rPr = run._element.get_or_add_rPr()
    # تثبيت الخط العربي
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), FONT_NAME)
    # تثبيت الحجم العربي
    szCs = docx.oxml.shared.OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(FONT_SIZE * 2))
    rPr.append(szCs)
    # تثبيت العرض العربي
    bCs = docx.oxml.shared.OxmlElement('w:bCs')
    bCs.set(qn('w:val'), '1')
    rPr.append(bCs)
    
    if is_italic:
        iCs = docx.oxml.shared.OxmlElement('w:iCs')
        iCs.set(qn('w:val'), '1')
        rPr.append(iCs)
    if is_red:
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(qn('w:val'), RED_HEX)
        rPr.append(color)

def process_transcript():
    if not os.path.exists(INPUT_FILE):
        print(f"خطأ: لم يتم العثور على الملف في المسار المذكور.")
        return

    print("جاري قراءة الملف ومعالجة انكسار الأسطر...")
    source_doc = docx.Document(INPUT_FILE)
    healed_sections = heal_broken_lines(source_doc)
    
    new_doc = docx.Document()
    # نمط البحث عن المتن المائل والآيات والأقواس
    pattern = re.compile(r'(\*.*?\*|﴿.*?﴾|\[\[.*?\]\]|\[.*?\]|\(.*?\))', re.DOTALL)

    for section_text in healed_sections:
        p = new_doc.add_paragraph()
        parts = pattern.split(section_text)
        
        for part in parts:
            if not part: continue
            
            # تحديد نوع النص بناءً على العلامات
            is_matn = part.startswith('*') and part.endswith('*')
            is_bracket = any(part.startswith(c) for c in ['﴿', '[', '('])
            
            # تنظيف النص من علامات الماركداون للعرض النهائي
            clean_text = part.replace('*', '')
            
            # منطق التلوين
            should_color = False
            if is_matn:
                should_color = True
            elif is_bracket and COLOR_BRACKETS:
                should_color = True
            
            run = p.add_run(clean_text)
            apply_complex_style(run, is_red=should_color, is_italic=is_matn)

    # حفظ ملف مؤقت تمهيداً لضبط الاتجاه عبر Word
    temp_path = OUTPUT_FILE.replace(".docx", "_temp.docx")
    new_doc.save(temp_path)
    
    try:
        print("جاري ضبط محاذاة النص من اليمين إلى اليسار...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        wdoc = word.Documents.Open(temp_path)
        
        # ضبط الفقرات لتكون RTL ومحاذاة لليمين
        wdoc.Content.ParagraphFormat.ReadingOrder = 1 # Right-to-Left
        wdoc.Content.ParagraphFormat.Alignment = 2    # Right Alignment
        
        wdoc.SaveAs(OUTPUT_FILE)
        wdoc.Close()
        word.Quit()
        if os.path.exists(temp_path):
            os.remove(temp_path)
        print(f"تم بنجاح! الملف النهائي: {os.path.basename(OUTPUT_FILE)}")
    except Exception as e:
        print(f"حدث خطأ أثناء ضبط المحاذاة: {e}")

if __name__ == "__main__":
    process_transcript()