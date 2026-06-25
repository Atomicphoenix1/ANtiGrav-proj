import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Pt, RGBColor
from fasils import TYPE_A_XML, TYPE_B_XML
import re
import sys
import os
import requests
import win32com.client
import time

# --- Configuration ---
N8N_WEBHOOK_URL = "https://atomicphoenix1.app.n8n.cloud/webhook/0b2e5c27-0dc5-4512-8fcb-cbec9ba785fa"

def apply_formatting(run, font_name, size, bold=False, color=None):
    """
    Surgically builds the Run Properties (rPr) to ensure both standard 
    and Complex Script (Arabic) properties are set in the correct XML order.
    """
    rPr = run._element.get_or_add_rPr()
    
    # Clear existing properties to ensure strict ordering
    for child in rPr.getchildren():
        rPr.remove(child)
        
    # 1. Fonts (Standard and Complex Script)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    
    # 2. Bold (Standard and Complex Script)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
        
    # 3. Color
    if color:
        c = OxmlElement('w:color')
        hex_color = "%02x%02x%02x" % (color[0], color[1], color[2])
        c.set(qn('w:val'), hex_color)
        rPr.append(c)
        
    # 4. Size (Standard and Complex Script) — value is in half-points
    sz_val = str(int(size * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    
    # 5. Right-to-Left marker (Crucial for Arabic punctuation positioning)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def insert_fasil(doc, xml_string):
    para = doc.add_paragraph()
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    element = parse_xml(xml_string)
    para._p.append(element)

def apply_rtl_justify(para):
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def process_paragraph_text(para, text, base_font, base_size, base_color=None):
    pattern = r'(<hadith>.*?</hadith>|<quran>.*?</quran>|<strong>.*?</strong>|\[.*?\]|صلى الله عليه وسلم|عز وجل|سبحانه وتعالى|تعالى)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part: continue
        
        run = para.add_run()
        
        if part.startswith('<hadith>') and part.endswith('</hadith>'):
            content = part.replace('<hadith>', '').replace('</hadith>', '')
            is_bold = False
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
                is_bold = True
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=is_bold, color=base_color)
            
        elif part.startswith('<quran>') and part.endswith('</quran>'):
            content = part.replace('<quran>', '').replace('</quran>', '')
            # If the verse starts and ends with decorative brackets
            if content.startswith('﴿') and content.endswith('﴾'):
                # 1. Opening bracket in Lotus font
                run.text = '﴿'
                apply_formatting(run, base_font, 18, color=RGBColor(0x20, 0x21, 0x22))
                # 2. Middle text in Uthmanic font
                run_mid = para.add_run(content[1:-1])
                apply_formatting(run_mid, 'KFGQPC Uthmanic Script HAFS', 18, color=RGBColor(0x20, 0x21, 0x22))
                # 3. Closing bracket in Lotus font
                run_end = para.add_run('﴾')
                apply_formatting(run_end, base_font, 18, color=RGBColor(0x20, 0x21, 0x22))
            else:
                run.text = content
                apply_formatting(run, 'KFGQPC Uthmanic Script HAFS', 18, color=RGBColor(0x20, 0x21, 0x22))
            
        elif part.startswith('<strong>') and part.endswith('</strong>'):
            content = part.replace('<strong>', '').replace('</strong>', '')
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=True, color=base_color)
            
        elif part.startswith('[') and part.endswith(']'):
            # Plain brackets = book titles or scholar quotes (NOT bold)
            content = part[1:-1]
            run.text = f">{content}<"
            apply_formatting(run, base_font, base_size, bold=False, color=base_color)
            
        elif part in ['صلى الله عليه وسلم', 'عز وجل', 'سبحانه وتعالى', 'تعالى']:
            run.text = part
            apply_formatting(run, 'DecoType Thuluth II', 30, bold=False, color=RGBColor(0xC0, 0x00, 0x00))
            
        else:
            run.text = part
            apply_formatting(run, base_font, base_size, bold=False, color=base_color)

def convert_to_pdf(docx_path, pdf_path):
    """Converts a Word document to PDF using win32com."""
    word = None
    wdoc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        
        wdoc = word.Documents.Open(abs_docx)
        # 17 = wdExportFormatPDF
        wdoc.ExportAsFixedFormat(abs_pdf, 17)
        return True
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return False
    finally:
        if wdoc: wdoc.Close(False)
        if word: word.Quit()

def send_to_n8n(docx_path, pdf_path):
    """Sends DOCX and PDF to n8n webhook."""
    try:
        files = {
            'docx': (os.path.basename(docx_path), open(docx_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'pdf': (os.path.basename(pdf_path), open(pdf_path, 'rb'), 'application/pdf'),
        }
        print(f"Sending to n8n: {os.path.basename(docx_path)}")
        resp = requests.post(N8N_WEBHOOK_URL, files=files)
        for f in files.values(): f[1].close()
        return resp.status_code in [200, 201]
    except Exception as e:
        print(f"n8n Upload Error: {e}")
        return False

def format_document(markdown_content, template_path, output_docx_path, is_file=True):
    """
    markdown_content: either a file path or a string of markdown.
    is_file: True if markdown_content is a path, False if it's raw text.
    """
    doc = docx.Document(template_path)
    
    # Fresh start
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
        
    if is_file:
        with open(markdown_content, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    else:
        lines = markdown_content.splitlines()
        
    for line in lines:
        line = line.strip()
        if not line: continue
            
        if line == '<matn>' or line == '</matn>': continue
            
        if line.startswith('<speaker>'):
            speaker_text = line.replace('<speaker>', '').replace('</speaker>', '')
            if 'قال الشارح' in speaker_text or 'قال المصنف' in speaker_text:
                if len(doc.paragraphs) > 0:
                    fasil_xml = TYPE_A_XML if 'قال الشارح' in speaker_text else TYPE_B_XML
                    insert_fasil(doc, fasil_xml)
            
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            process_paragraph_text(para, speaker_text, 'DecoType Thuluth II', 30, RGBColor(0xC0, 0x00, 0x00))
            continue
            
        if line.startswith('- '):
            line = line[2:]
            para = doc.add_paragraph(style='List Paragraph')
            apply_rtl_justify(para)
            bullet_run = para.add_run("○ ")
            apply_formatting(bullet_run, 'AAAGoldenLotus Stg1_Ver1', 18, color=RGBColor(0x00, 0xB0, 0x50))
            process_paragraph_text(para, line, 'AAAGoldenLotus Stg1_Ver1', 18)
            continue
            
        para = doc.add_paragraph()
        apply_rtl_justify(para)
        process_paragraph_text(para, line, 'AAAGoldenLotus Stg1_Ver1', 18)
    
    # Final Green Symbol
    insert_fasil(doc, TYPE_A_XML)
    
    doc.save(output_docx_path)
    print(f"Saved DOCX: {output_docx_path}")
    
    # PDF Conversion
    pdf_path = output_docx_path.replace(".docx", ".pdf")
    if convert_to_pdf(output_docx_path, pdf_path):
        print(f"Saved PDF: {pdf_path}")
        # Send to n8n
        if send_to_n8n(output_docx_path, pdf_path):
            print("Successfully sent to n8n/Telegram!")
    
    return output_docx_path, pdf_path

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(base_dir)
    md_path = os.path.join(base_dir, "sample.md")
    template_path = os.path.join(parent_dir, "template.docx")
    output_path = os.path.join(base_dir, "output.docx")
    format_document(md_path, template_path, output_path)
