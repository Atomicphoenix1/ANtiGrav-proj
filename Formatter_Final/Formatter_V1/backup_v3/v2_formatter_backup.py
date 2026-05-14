import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Pt, RGBColor
from fasils_backup  import TYPE_A_XML, TYPE_B_XML
import re
import sys
import os

def apply_formatting(run, font_name, size, bold=False, color=None):
    """
    Surgically builds the Run Properties (rPr) to ensure both standard 
    and Complex Script (Arabic) properties are set in the correct XML order.
    """
    rPr = run._element.get_or_add_rPr()
    
    # Clear existing properties to ensure strict ordering
    for child in rPr.getchildren():
        rPr.remove(child)
        
    # 1. Fonts (Standard and Complex Script)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    
    # 2. Bold (Standard and Complex Script)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
        
    # 3. Color
    if color:
        c = OxmlElement('w:color')
        # Convert RGBColor to hex string
        hex_color = "%02x%02x%02x" % (color[0], color[1], color[2])
        c.set(qn('w:val'), hex_color)
        rPr.append(c)
        
    # 4. Size (Standard and Complex Script) - value is in half-points
    sz_val = str(int(size * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    
    # 5. Right-to-Left marker (Crucial for punctuation)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def insert_fasil(doc, xml_string):
    para = doc.add_paragraph()
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    element = parse_xml(xml_string)
    para._p.append(element)

def apply_rtl_justify(para):
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def process_paragraph_text(para, text, base_font, base_size, base_color=None):
    pattern = r'(<hadith>.*?</hadith>|<quran>.*?</quran>|<strong>.*?</strong>|\[.*?\]|صلى الله عليه وسلم|عز وجل|سبحانه وتعالى|تعالى)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part: continue
        
        run = para.add_run()
        
        if part.startswith('<hadith>') and part.endswith('</hadith>'):
            content = part.replace('<hadith>', '').replace('</hadith>', '')
            is_bold = False
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
                is_bold = True
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=is_bold, color=base_color)
            
        elif part.startswith('<quran>') and part.endswith('</quran>'):
            content = part.replace('<quran>', '').replace('</quran>', '')
            is_bold = False
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
                is_bold = True
            else:
                run.text = content
            apply_formatting(run, 'KFGQPC Uthmanic Script HAFS', 18, bold=is_bold, color=RGBColor(0x20, 0x21, 0x22))
            
        elif part.startswith('<strong>') and part.endswith('</strong>'):
            content = part.replace('<strong>', '').replace('</strong>', '')
            if content.startswith('[') and content.endswith(']'):
                run.text = f">{content[1:-1]}<"
            else:
                run.text = content
            apply_formatting(run, base_font, base_size, bold=True, color=base_color)
            
        elif part.startswith('[') and part.endswith(']'):
            content = part[1:-1]
            run.text = f">{content}<"
            apply_formatting(run, base_font, base_size, bold=True, color=base_color)
            
        elif part in ['صلى الله عليه وسلم', 'عز وجل', 'سبحانه وتعالى', 'تعالى']:
            run.text = part
            apply_formatting(run, 'DecoType Thuluth II', 30, bold=False, color=RGBColor(0xC0, 0x00, 0x00))
            
        else:
            run.text = part
            apply_formatting(run, base_font, base_size, bold=False, color=base_color)

def format_document(markdown_path, template_path, output_path):
    doc = docx.Document(template_path)
    
    # Delete all existing paragraphs in the template to start fresh
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None
        
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_matn = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line == '<matn>':
            in_matn = True
            continue
        elif line == '</matn>':
            in_matn = False
            continue
            
        if line.startswith('<matn>') and line.endswith('</matn>'):
            line = line.replace('<matn>', '').replace('</matn>', '')
            
        # Detect Speaker
        if line.startswith('<speaker>'):
            speaker_text = line.replace('<speaker>', '').replace('</speaker>', '')
            
            if 'قال الشارح' in speaker_text:
                insert_fasil(doc, TYPE_A_XML)
            elif 'قال المصنف' in speaker_text:
                if len(doc.paragraphs) > 0:
                    insert_fasil(doc, TYPE_B_XML)
                
            para = doc.add_paragraph()
            # Apply RTL to paragraph
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            process_paragraph_text(para, speaker_text, 'DecoType Thuluth II', 30, RGBColor(0xC0, 0x00, 0x00))
            continue
            
        # Detect Lists
        if line.startswith('- '):
            line = line[2:]
            para = doc.add_paragraph(style='List Paragraph')
            apply_rtl_justify(para)
            process_paragraph_text(para, line, 'AAAGoldenLotus Stg1_Ver1', 18)
            continue
            
        # Standard Paragraph
        para = doc.add_paragraph()
        apply_rtl_justify(para)
        
        base_font = 'AAAGoldenLotus Stg1_Ver1'
        base_size = 18
        
        process_paragraph_text(para, line, base_font, base_size)
        
    doc.save(output_path)
    print(f"Saved formatted document to: {output_path}")

if __name__ == "__main__":
    base_dir = r"C:\Users\saif_\Desktop\downs\حاليًا\يومي\Lectures\ANtiGrav"
    markdown_path = os.path.join(base_dir, "Formatter_V2", "sample.md")
    template_path = os.path.join(base_dir, "تفريغ كتاب التوحيد - باب 9.docx")
    output_path = os.path.join(base_dir, "Formatter_V2", "Final_Formatted.docx")
    
    sys.path.append(os.path.join(base_dir, "Formatter_V2"))
    
    format_document(markdown_path, template_path, output_path)
